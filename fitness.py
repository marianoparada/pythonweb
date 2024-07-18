import streamlit as st
import random
import time
from docx import Document
from docx.shared import Inches
import io

def guardar_rutina_word(rutina):
    doc = Document()
    doc.add_heading('FA FITNESS - Rutina Personalizada', 0)

    vuelta_actual = 1
    for grupo, ejercicio, duracion, instruccion, vuelta, num_ejercicio, total_ejercicios in rutina:
        if vuelta != vuelta_actual:
            doc.add_paragraph('---')
            doc.add_paragraph(f'Vuelta {vuelta}')
            vuelta_actual = vuelta
        if grupo == "Descanso":
            if ejercicio == "Descanso entre vueltas":
                doc.add_paragraph(f'Descanso entre vueltas: {duracion} segundos')
            else:
                doc.add_paragraph(f'Descanso: {duracion} segundos')
        else:
            doc.add_paragraph(f'Ejercicio {num_ejercicio}/{total_ejercicios}: {grupo} - {ejercicio}: {duracion} segundos')
            doc.add_paragraph(f'Instrucción: {instruccion}')

    # Guardar el documento en un objeto BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Botón de descarga
    st.download_button(
        label="Guardar Rutina",
        data=buffer,
        file_name="FA_FITNESS_Rutina.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="descargar_rutina"
    )
# Configuración de la página
st.set_page_config(page_title="FA FITNESS - Generador de Rutinas", layout="wide")


ejercicios = {
    "💪 Brazos": [
        ("Curl de bíceps (mancuernas)", "De pie, sostenga una mancuerna en cada mano con los brazos extendidos. Doble los codos para levantar las mancuernas hacia los hombros, luego baje lentamente."),
        ("Extensiones de tríceps (mancuernas)", "Siéntese o póngase de pie con una mancuerna sostenida con ambas manos sobre la cabeza. Baje la mancuerna detrás de la cabeza doblando los codos, luego extienda los brazos."),
        ("Flexiones de brazos", "Acuéstese boca abajo con las manos apoyadas en el suelo al ancho de los hombros. Empuje el cuerpo hacia arriba manteniendo el cuerpo recto, luego baje."),
        ("Fondos de tríceps (silla)", "Siéntese en el borde de una silla, coloque las manos a los lados y deslice el cuerpo hacia adelante. Baje el cuerpo doblando los codos y luego empuje hacia arriba."),
        ("Curl martillo (mancuernas)", "De pie, sostenga una mancuerna en cada mano con las palmas hacia el cuerpo. Doble los codos para levantar las mancuernas hacia los hombros, luego baje lentamente."),
        ("Flexiones diamante", "Acuéstese boca abajo con las manos juntas formando un diamante debajo del pecho. Empuje el cuerpo hacia arriba manteniendo el cuerpo recto, luego baje."),
        ("Curl de bíceps concentrado (mancuernas)", "Siéntese en un banco, sostenga una mancuerna en una mano y apoye el codo en el muslo. Doble el codo para levantar la mancuerna, luego baje lentamente."),
        ("Extensiones de tríceps con cuerda (máquina)", "De pie, sostenga la cuerda conectada a la máquina de polea alta con ambas manos. Baje la cuerda doblando los codos, luego extienda los brazos."),
        ("Curl de bíceps en banco inclinado (mancuernas)", "Siéntese en un banco inclinado, sostenga una mancuerna en cada mano con los brazos colgando hacia abajo. Doble los codos para levantar las mancuernas hacia los hombros, luego baje lentamente."),
        ("Extensiones de tríceps sobre la cabeza (mancuernas)", "Siéntese o póngase de pie con una mancuerna sostenida con ambas manos sobre la cabeza. Baje la mancuerna detrás de la cabeza doblando los codos, luego extienda los brazos.")
    ],
    "🦵 Piernas": [
        ("Sentadillas", "De pie con los pies separados al ancho de los hombros, baje el cuerpo como si fuera a sentarse en una silla invisible, manteniendo el pecho erguido. Luego, vuelva a la posición inicial."),
        ("Estocadas", "De pie, dé un paso largo hacia adelante con una pierna. Baje el cuerpo hasta que ambas rodillas estén dobladas en ángulos de 90 grados. Empuje hacia atrás para volver a la posición inicial y alterne las piernas."),
        ("Peso muerto (barra)", "De pie con los pies separados al ancho de los hombros, sostenga una barra frente a los muslos. Inclínese hacia adelante desde las caderas, manteniendo la espalda recta, hasta que la barra llegue a las espinillas. Luego, vuelva a la posición inicial."),
        ("Elevaciones de talones", "De pie, levántese sobre los dedos de los pies, manteniendo las rodillas rectas. Baje lentamente los talones hacia el suelo."),
        ("Puente de glúteos", "Acuéstese boca arriba con las rodillas dobladas y los pies apoyados en el suelo. Levante las caderas hacia el techo, apretando los glúteos, luego baje lentamente."),
        ("Prensa de pierna (máquina)", "Siéntese en la máquina de prensa de pierna con los pies apoyados en la plataforma. Empuje la plataforma hacia adelante hasta que las piernas estén extendidas, luego baje lentamente."),
        ("Extensiones de pierna (máquina)", "Siéntese en la máquina de extensión de pierna con las piernas debajo de la almohadilla. Extienda las piernas hacia adelante, luego baje lentamente."),
        ("Curl de pierna (máquina)", "Acuéstese boca abajo en la máquina de curl de pierna con los tobillos debajo de la almohadilla. Doble las rodillas para levantar la almohadilla hacia los glúteos, luego baje lentamente."),
        ("Sentadillas búlgaras (mancuernas)", "De pie, coloque un pie en un banco detrás de usted y sostenga una mancuerna en cada mano. Baje el cuerpo doblando la rodilla de la pierna delantera, luego empuje hacia arriba."),
        ("Saltos al cajón", "De pie frente a un cajón, salte con ambos pies para aterrizar en el cajón, luego baje de un paso.")
    ],
    "🏋️ Abdominales": [
        ("Abdominales", "Acuéstese boca arriba con las rodillas dobladas y los pies apoyados en el suelo. Levante la parte superior del cuerpo hacia las rodillas, luego baje lentamente."),
        ("Plancha", "Acuéstese boca abajo con los antebrazos apoyados en el suelo y los codos debajo de los hombros. Levante el cuerpo manteniéndolo recto desde la cabeza hasta los pies."),
        ("Giros rusos", "Siéntese en el suelo con las rodillas dobladas y los pies elevados. Sostenga un peso con ambas manos y gire el torso de un lado a otro."),
        ("Elevaciones de piernas", "Acuéstese boca arriba con las piernas rectas. Levante las piernas hacia el techo hasta que los glúteos se despeguen del suelo, luego baje lentamente."),
        ("Bicicleta", "Acuéstese boca arriba con las manos detrás de la cabeza y las piernas levantadas. Alterne llevando el codo hacia la rodilla opuesta mientras extiende la otra pierna."),
        ("Plancha lateral", "Acuéstese de lado con un antebrazo apoyado en el suelo y el codo debajo del hombro. Levante el cuerpo manteniéndolo recto desde la cabeza hasta los pies."),
        ("Crunch inverso", "Acuéstese boca arriba con las piernas dobladas y los pies elevados. Levante las caderas hacia el techo, luego baje lentamente."),
        ("Escaladores", "Empiece en posición de flexión de brazos. Lleve una rodilla hacia el pecho, luego alterne rápidamente las piernas."),
        ("V-ups", "Acuéstese boca arriba con los brazos extendidos sobre la cabeza. Levante simultáneamente las piernas y el torso para tocarse los pies, luego baje lentamente."),
        ("Plancha con elevación de brazo", "Empiece en posición de plancha. Levante un brazo extendido hacia adelante, luego alterne los brazos.")
    ],
    "🏋️ Hombros": [
        ("Press militar (mancuernas)", "De pie o sentado, sostenga una mancuerna en cada mano a la altura de los hombros. Empuje las mancuernas hacia arriba hasta que los brazos estén extendidos, luego baje lentamente."),
        ("Elevaciones laterales (mancuernas)", "De pie, sostenga una mancuerna en cada mano a los lados. Levante los brazos hacia los lados hasta que estén a la altura de los hombros, luego baje lentamente."),
        ("Face pulls (cable)", "De pie, sostenga la cuerda conectada a la máquina de polea alta con ambas manos. Tire de la cuerda hacia la cara, manteniendo los codos altos."),
        ("Press Arnold (mancuernas)", "Sentado, sostenga una mancuerna en cada mano a la altura de los hombros con las palmas hacia adentro. Gire las palmas hacia afuera mientras empuja las mancuernas hacia arriba, luego baje lentamente."),
        ("Elevaciones frontales (mancuernas)", "De pie, sostenga una mancuerna en cada mano frente a los muslos. Levante los brazos hacia adelante hasta que estén a la altura de los hombros, luego baje lentamente."),
        ("Remo al mentón (barra)", "De pie, sostenga una barra con las manos juntas frente a los muslos. Levante la barra hacia el mentón, manteniendo los codos altos, luego baje lentamente."),
        ("Encogimientos de hombros (mancuernas)", "De pie, sostenga una mancuerna en cada mano a los lados. Levante los hombros hacia las orejas, luego baje lentamente."),
        ("Elevaciones posteriores (mancuernas)", "De pie, inclínese hacia adelante desde las caderas con una mancuerna en cada mano. Levante los brazos hacia los lados hasta que estén a la altura de los hombros, luego baje lentamente."),
        ("Press de hombros con barra", "Sentado o de pie, sostenga una barra a la altura de los hombros con las palmas hacia adelante. Empuje la barra hacia arriba hasta que los brazos estén extendidos, luego baje lentamente."),
        ("Remo con barra en T", "De pie, inclínese hacia adelante desde las caderas y sostenga una barra en T con ambas manos. Tire de la barra hacia el pecho, luego baje lentamente.")
    ],
    "🏃 Aeróbico": [
        ("Saltar la soga", "Salte con ambos pies mientras gira la soga por encima y por debajo del cuerpo."),
        ("Burpees", "Desde una posición de pie, agáchese y coloque las manos en el suelo. Salte los pies hacia atrás para llegar a una posición de flexión, haga una flexión, salte los pies hacia adelante y levántese saltando."),
        ("Escaladores", "Empiece en posición de flexión de brazos. Lleve una rodilla hacia el pecho, luego alterne rápidamente las piernas."),
        ("Correr", "Corra a un ritmo constante durante un período de tiempo o distancia."),
        ("Piques", "Corra a máxima velocidad durante una distancia corta, luego descanse y repita."),
        ("Saltos de tijera", "Desde una posición de pie, salte abriendo las piernas y levantando los brazos por encima de la cabeza, luego vuelva a la posición inicial."),
        ("Rodillas altas", "Corra en el lugar llevando las rodillas lo más alto posible."),
        ("Trote en el lugar", "Corra suavemente en el lugar, levantando los pies del suelo."),
        ("Escaladores", "Empiece en posición de flexión de brazos. Lleve una rodilla hacia el pecho, luego alterne rápidamente las piernas."),
        ("Saltos al banco", "De pie frente a un banco, salte con ambos pies para aterrizar en el banco, luego baje de un paso.")
    ]
}

# Distribución de ejercicios según la priorización (sin cambios)
distribucion_ejercicios = {
    "Tren superior": {"💪 Brazos": 3, "🏋️ Hombros": 2, "🏋️ Abdominales": 1, "🦵 Piernas": 1, "🏃 Aeróbico": 1},
    "Zona media": {"🏋️ Abdominales": 3, "💪 Brazos": 2, "🦵 Piernas": 1, "🏋️ Hombros": 1, "🏃 Aeróbico": 1},
    "Tren inferior": {"🦵 Piernas": 3, "🏋️ Abdominales": 2, "💪 Brazos": 1, "🏋️ Hombros": 1, "🏃 Aeróbico": 1},
    "Aeróbico": {"🏃 Aeróbico": 3, "🦵 Piernas": 2, "💪 Brazos": 1, "🏋️ Abdominales": 1, "🏋️ Hombros": 1}
}

def generar_rutina(prioridad, duracion, tiempo_descanso, vueltas):
    rutina_base = []
    for grupo, cantidad in distribucion_ejercicios[prioridad].items():
        ejercicios_grupo = random.sample(ejercicios[grupo], cantidad)
        for ejercicio, instruccion in ejercicios_grupo:
            if duracion == "Ambos aleatorios":
                tiempo = random.choice([30, 40])
            else:
                tiempo = int(duracion.split()[0])
            rutina_base.append((grupo, ejercicio, tiempo, instruccion))
    
    rutina_completa = []
    for vuelta in range(1, vueltas + 1):
        for i, ejercicio in enumerate(rutina_base, 1):
            grupo, nombre_ejercicio, tiempo, instruccion = ejercicio
            rutina_completa.append((grupo, nombre_ejercicio, tiempo, instruccion, vuelta, i, len(rutina_base)))
            if i < len(rutina_base):  # Añadir descanso después de cada ejercicio, excepto el último
                rutina_completa.append(("Descanso", "Descanso entre ejercicios", tiempo_descanso, "Toma un breve descanso antes del siguiente ejercicio", vuelta, None, None))
        if vuelta < vueltas:  # Añadir descanso largo entre vueltas
            rutina_completa.append(("Descanso", "Descanso entre vueltas", tiempo_descanso * 5, "Toma un descanso más largo antes de la siguiente vuelta", vuelta, None, None))
    
    return rutina_completa

def mostrar_rutina(rutina):
    st.title("🏋️ FA FITNESS - Rutina Personalizada")
    vuelta_actual = 1
    for grupo, ejercicio, duracion, instruccion, vuelta, num_ejercicio, total_ejercicios in rutina:
        if vuelta != vuelta_actual:
            st.write("---")
            st.write(f"Vuelta {vuelta}")
            vuelta_actual = vuelta
        if grupo == "Descanso":
            if ejercicio == "Descanso entre vueltas":
                st.write(f"Descanso entre vueltas: {duracion} segundos")
            else:
                st.write(f"Descanso: {duracion} segundos")
        else:
            st.write(f"Ejercicio {num_ejercicio}/{total_ejercicios}: {grupo} - {ejercicio}: {duracion} segundos")
            st.write(f"Instrucción: {instruccion}")

def temporizador(rutina):
    st.title("🏋️ FA FITNESS - Temporizador de Rutina")
    
    if 'ejercicio_actual' not in st.session_state:
        st.session_state.ejercicio_actual = 0
        st.session_state.tiempo_restante = rutina[0][2]
    
    grupo, ejercicio, duracion, instruccion, vuelta, num_ejercicio, total_ejercicios = rutina[st.session_state.ejercicio_actual]
    
    col1, col2, col3 = st.columns([1,2,1])
    with col2:
        if grupo != "Descanso":
            st.markdown(f"<h3 style='text-align: center;'>Vuelta {vuelta}/{st.session_state.vueltas} - Ejercicio {num_ejercicio}/{total_ejercicios}</h3>", unsafe_allow_html=True)
        
        if grupo == "Descanso":
            if ejercicio == "Descanso entre vueltas":
                st.markdown("<h1 style='text-align: center;'>Descanso entre vueltas</h1>", unsafe_allow_html=True)
            else:
                st.markdown("<h1 style='text-align: center;'>Descanso</h1>", unsafe_allow_html=True)
            # Mostrar el próximo ejercicio
            proximo_indice = st.session_state.ejercicio_actual + 1
            if proximo_indice < len(rutina):
                proximo_ejercicio = rutina[proximo_indice][1]
                st.markdown(f"<h3 style='text-align: center;'>Próximo ejercicio: {proximo_ejercicio}</h3>", unsafe_allow_html=True)
        else:
            st.markdown(f"<h1 style='text-align: center;'>{grupo}</h1>", unsafe_allow_html=True)
            st.markdown(f"<h2 style='text-align: center;'>{ejercicio}</h2>", unsafe_allow_html=True)
            st.markdown(f"<p style='text-align: center;'>{instruccion}</p>", unsafe_allow_html=True)
        
        progress = st.progress(0)
        tiempo_texto = st.empty()
        
        if st.button("Pausar/Reanudar", key="pausar_reanudar"):
            st.session_state.timer_running = not st.session_state.timer_running
    
    if st.session_state.timer_running:
        progress.progress(1 - st.session_state.tiempo_restante / duracion)
        tiempo_texto.markdown(f"<h3 style='text-align: center;'>{st.session_state.tiempo_restante} segundos</h3>", unsafe_allow_html=True)
        
        time.sleep(1)
        st.session_state.tiempo_restante -= 1
        
        if st.session_state.tiempo_restante < 0:
            st.session_state.ejercicio_actual += 1
            if st.session_state.ejercicio_actual < len(rutina):
                st.session_state.tiempo_restante = rutina[st.session_state.ejercicio_actual][2]
            else:
                st.success("¡Felicidades! Has completado tu rutina.")
                if st.button("Volver al inicio", key="volver_inicio"):
                    for key in list(st.session_state.keys()):
                        del st.session_state[key]
                    st.experimental_rerun()
                return
        
        st.experimental_rerun()

def main():
    if 'rutina' not in st.session_state:
        st.title("🏋️ FA FITNESS - Generador de Rutinas")
        
        prioridad = st.selectbox(
            "Selecciona la prioridad de tu rutina:",
            ["Tren superior", "Zona media", "Tren inferior", "Aeróbico"]
        )
        
        duracion = st.selectbox(
            "Selecciona la duración de los ejercicios:",
            ["30 segundos", "40 segundos", "Ambos aleatorios"]
        )
        
        tiempo_descanso = st.number_input(
            "Selecciona el tiempo de descanso entre ejercicios (en segundos):",
            min_value=5,
            max_value=60,
            value=15,
            step=5
        )
        
        vueltas = st.number_input(
            "Número de vueltas:",
            min_value=1,
            max_value=10,
            value=1
        )
        
        if st.button("Generar Rutina", key="generar_rutina"):
            with st.spinner("Generando tu rutina personalizada..."):
                time.sleep(2)  # Espera 2 segundos
                st.session_state.rutina = generar_rutina(prioridad, duracion, tiempo_descanso, vueltas)
                st.session_state.prioridad = prioridad
                st.session_state.duracion = duracion
                st.session_state.tiempo_descanso = tiempo_descanso
                st.session_state.vueltas = vueltas
            st.success("¡Rutina generada con éxito!")
            st.experimental_rerun()
    
    elif 'timer_running' not in st.session_state:
        col1, col2, col3 = st.columns([1,2,1])
        with col2:
            st.markdown(
                """
                <style>
                div.stButton > button {
                    width: 100%;
                    height: 3em;
                    font-size: 20px;
                }
                </style>
                """, 
                unsafe_allow_html=True
            )
            if st.button("Comenzar Rutina", key="comenzar_rutina"):
                st.session_state.timer_running = True
                st.session_state.ejercicio_actual = 0
                st.session_state.tiempo_restante = st.session_state.rutina[0][2]
                st.experimental_rerun()
            
            guardar_rutina_word(st.session_state.rutina)
    
    else:
        temporizador(st.session_state.rutina)

if __name__ == "__main__":
    main()